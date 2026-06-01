"""
Generates TARA_SOW.docx from structured content.
Run from the restiview project root:
    python tool/make_sow.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# â"€â"€ Brand colours â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
DARK_GREEN  = RGBColor(0x2E, 0x4F, 0x3E)
BEIGE_BG    = RGBColor(0xF5, 0xF0, 0xE6)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF0, 0xF0, 0xF0)
MID_GREY    = RGBColor(0x60, 0x60, 0x60)
BLACK       = RGBColor(0x00, 0x00, 0x00)

OUTPUT_PATH = "TARA_SOW.docx"


# â"€â"€ Helpers â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_col = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    tcPr.append(shd)


def add_cover(doc: Document):
    # Large top margin space
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RestiView")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = DARK_GREEN

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Design & Launch Support")
    run2.font.size = Pt(22)
    run2.font.color.rgb = DARK_GREEN

    doc.add_paragraph()

    sow = doc.add_paragraph()
    sow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sow.add_run("Statement of Work")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = MID_GREY

    doc.add_paragraph()

    meta_lines = [
        ("Client",     "Peter Kern"),
        ("Contractor", "Tara Kern"),
        ("Date",       "May 2026"),
        ("Version",    "1.0"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}:  ")
        rl.font.bold = True
        rl.font.size = Pt(11)
        rv = p.add_run(value)
        rv.font.size = Pt(11)

    doc.add_page_break()


def h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = WHITE
    # shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2E4F3E')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Cm(0.3)


def h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_GREEN


def h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = MID_GREY


def body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)


def label_value(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rl = p.add_run(f"{label}:  ")
    rl.font.bold = True
    rl.font.size = Pt(10.5)
    rv = p.add_run(value)
    rv.font.size = Pt(10.5)


def divider(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E4F3E')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(doc: Document, headers: list, rows: list):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
        set_cell_bg(hdr_cells[i], DARK_GREEN)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = LIGHT_GREY if ri % 2 == 0 else WHITE
        for ci, cell_text in enumerate(row_data):
            cells[ci].text = cell_text
            set_cell_bg(cells[ci], bg)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table


def add_task(doc, ref, title, description_paras, bullets_label, bullets, output_label, outputs):
    h2(doc, f"{ref} â€" {title}")
    h3(doc, "Description")
    for para in description_paras:
        body(doc, para)
    if bullets:
        body(doc, bullets_label)
        for b in bullets:
            bullet(doc, b)
    h3(doc, "Expected Output")
    for o in outputs:
        bullet(doc, o)
    divider(doc)


# â"€â"€ Document assembly â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# â"€â"€ Cover page â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
add_cover(doc)

# â"€â"€ Section 1: Overview â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "1. Overview")
body(doc,
    "Tara Kern ("Contractor") is engaged to provide graphic design, UX review, testing, "
    "and marketing asset production services to support the public launch of the RestiView "
    "mobile application on Google Play.")
body(doc,
    "RestiView is a restaurant review app for Android. Users can record and manage personal "
    "restaurant reviews, share reviews with friends, and discover nearby restaurants. The app "
    "is currently in closed beta testing and is being prepared for public release on the "
    "Google Play Store.")

# â"€â"€ Section 2: Scope â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "2. Scope of Work")

# Task A
h1(doc, "Task A â€" App Design Review & Improvements")

add_task(doc, "A1", "Full Visual & UX Audit",
    ["Review every screen in the RestiView app and document current design issues. This includes "
     "assessing colour usage, typography consistency, spacing, layout balance, button sizing, "
     "touch target sizes, and overall visual polish.",
     "The review should consider both aesthetic quality and usability â€" for example, whether labels "
     "are clear, whether actions are obvious, and whether the app feels professional and consistent "
     "throughout."],
    "", [],
    "Deliverable:",
    ["A written or annotated document (screenshots with notes, or a Figma/design tool file) listing "
     "each screen with identified issues and suggested improvements, prioritised by impact."]
)

add_task(doc, "A2", "Revised Colour Palette & Typography Proposal",
    ["Based on the audit findings, propose a revised visual direction for the app. The current "
     "colour scheme uses dark green (#2E4F3E) and beige (#F5F0E6) as primary brand colours. "
     "The proposal should either refine these or suggest an alternative palette. Typography "
     "should be assessed â€" the app currently uses the Gelica font family."],
    "The proposal must include:",
    ["Primary, secondary, and accent colours with hex values",
     "Background and surface colours",
     "Text colours (primary, secondary, muted)",
     "Font pairing recommendation (or confirmation to keep Gelica)",
     "Recommended font sizes for headings, body, labels, and buttons"],
    "Deliverable:",
    ["A colour and typography specification â€" PDF, Figma style guide page, or equivalent.",
     "Must include hex values for all colours so the developer can implement them directly in code.",
     "Source file (Figma, Illustrator, or equivalent)."]
)

add_task(doc, "A3", "Design Implementation Collaboration",
    ["Work with the developer to implement the agreed design changes in the app. The developer will "
     "make all code changes. The Contractor's role is to review the results on the device, identify "
     "where the implementation doesn't match the proposal, and provide clear feedback until the "
     "result matches the agreed design."],
    "", [],
    "Deliverable:",
    ["Written sign-off confirming the implemented design matches the approved proposal.",
     "Before/after screenshot comparison (optional but helpful)."]
)

add_task(doc, "A4", "Button & Component Consistency Review",
    ["Review the consistency of interactive elements across the app â€" buttons, input fields, "
     "dropdowns, dialogs, and navigation elements. Assess whether sizes, colours, and styles "
     "are consistent throughout.",
     "For example: are all primary action buttons the same colour? Are all destructive actions "
     "(delete, clear) visually distinct? Are touch targets large enough for comfortable use?"],
    "", [],
    "Deliverable:",
    ["A specification document or annotated screenshots defining the intended style for each "
     "component type: primary button, secondary button, destructive button, text button, "
     "input field, etc., so the developer can apply them consistently."]
)

# Task B
h1(doc, "Task B â€" App Testing")

add_task(doc, "B1", "Structured App Walkthrough",
    ["Test the app by working through all main user flows. For each flow, note anything that is "
     "confusing, broken, missing, or visually inconsistent."],
    "Flows to cover:",
    ["Registration and sign-in",
     "Adding a new restaurant review (general info, ratings, comments, photos)",
     "Editing and deleting a review",
     "Viewing the review list and filtering/sorting",
     "Preview screen",
     "Friends â€" sending and receiving friend requests",
     "Requesting and sharing reviews between friends",
     "Settings screen",
     "Help screen"],
    "Deliverable:",
    ["A bug and feedback log â€" spreadsheet or document with columns: Screen, Issue Description, "
     "Severity (High / Medium / Low), Suggested Fix."]
)

add_task(doc, "B2", "Retest After Fixes",
    ["After the developer has addressed the issues found in B1, retest the affected areas to "
     "confirm fixes are working correctly and no new issues have been introduced."],
    "", [],
    "Deliverable:",
    ["Updated bug log with each issue marked as Resolved, Partially Fixed, or Still Outstanding."]
)

# Task C
h1(doc, "Task C â€" Play Store Launch Assets")

add_task(doc, "C1", "App Icon",
    ["Design a polished app icon for the Google Play Store. The icon must work at small sizes "
     "(as small as 48Ã—48 on a device home screen) as well as large (512Ã—512 on the Play Store "
     "listing). It should be visually distinctive, recognisable, and reflect the restaurant "
     "review theme of the app.",
     "Note: do not apply rounded corners â€" Android applies these automatically."],
    "Requirements:",
    ["512Ã—512 PNG with transparent or solid background",
     "Adaptive icon foreground layer (safe zone: 108Ã—108dp)",
     "Adaptive icon background layer"],
    "Deliverables:",
    ["icon_512.png (512Ã—512 â€" Play Store upload)",
     "icon_foreground.png (adaptive icon foreground layer)",
     "icon_background.png (adaptive icon background layer)",
     "Source file (Figma, Illustrator, or equivalent)"]
)

add_task(doc, "C2", "Feature Graphic",
    ["Design the feature graphic shown at the top of the Play Store listing page. This is a "
     "banner-style image that introduces the app visually. It should include the app name "
     "'RestiView', a tagline or short description, and visual elements that communicate what "
     "the app does.",
     "Important: must not include device frames or screenshots â€" Google prohibits this in the "
     "feature graphic. Should look good on both light and dark Play Store themes."],
    "Requirements:",
    ["Size: 1024Ã—500 pixels", "Format: JPG or PNG"],
    "Deliverables:",
    ["feature_graphic.png (1024Ã—500)",
     "Source file"]
)

add_task(doc, "C3", "Short Description (Play Store)",
    ["Write the short description for the Play Store listing. This appears below the app name "
     "in search results and on the listing page. It must be punchy, clear, and within 80 "
     "characters. It should communicate what the app does and why someone would want it."],
    "", [],
    "Deliverable:",
    ["One or two alternative short description options (80 characters max each) for Peter to choose from."]
)

add_task(doc, "C4", "Full Description (Play Store)",
    ["Write the full app description for the Play Store listing (up to 4000 characters). "
     "This should explain what RestiView does, list key features, and be written in a way "
     "that is both engaging to a human reader and discoverable via Play Store search â€" "
     "naturally including relevant keywords such as 'restaurant review', 'dining', 'food "
     "diary', etc."],
    "Suggested structure:",
    ["Opening hook (1â€"2 sentences)",
     "What is RestiView? (short paragraph)",
     "Key features (bullet list)",
     "Who is it for? (short paragraph)",
     "Call to action"],
    "Deliverable:",
    ["Draft full description ready to paste into Play Console, in plain text.",
     "Include a note of the keywords used."]
)

# Task D
h1(doc, "Task D â€" Website Update")
body(doc,
    "Note: The RestiView website (restiview.com) is hosted on Yola and managed via their "
    "SiteBuilder tool. No coding is required â€" all changes are made through the visual editor. "
    "The Privacy Policy page is already complete and does not need updating.")
divider(doc)

add_task(doc, "D1", "Content & Design Audit",
    ["Review the current website and document what needs to be updated or improved. Consider: "
     "Is the content current and accurate? Does the visual design match the updated app design? "
     "Is the call-to-action (download the app) clear and prominent? Are there any missing pages?"],
    "", [],
    "Deliverable:",
    ["A short written list of recommended changes, prioritised."]
)

add_task(doc, "D2", "Website Content Update",
    ["Using the Yola SiteBuilder, update the website content to reflect the current state of "
     "the app."],
    "Updates to include:",
    ["Update text descriptions to match the current feature set",
     "Replace any outdated screenshots with current ones",
     "Add the Play Store download badge and link once the app is publicly live",
     "Ensure the app icon and feature graphic are used consistently across the site",
     "Add a Terms & Conditions page (required for Apple App Store submission in a future phase)"],
    "Deliverable:",
    ["Updated live website with all changes published at restiview.com."]
)

# Task E
h1(doc, "Task E â€" Launch & Marketing Materials")

add_task(doc, "E1", "Social Media Announcement Graphics",
    ["Design a set of graphics for announcing the app launch on social media. These should be "
     "visually consistent with the app and website, include the app name, a short message "
     "(e.g. 'Now available on Google Play'), and a clear call to action."],
    "Sizes required:",
    ["Square post: 1080Ã—1080 (Facebook, Instagram)",
     "Landscape: 1200Ã—628 (Facebook link share, X/Twitter)"],
    "Deliverables:",
    ["social_square.png (1080Ã—1080)",
     "social_landscape.png (1200Ã—628)",
     "Source files"]
)

add_task(doc, "E2", "Launch Email Template",
    ["Write and design a short email that Peter can send to his contacts announcing the app "
     "launch. Should include a brief description of the app, a screenshot or graphic, and a "
     "link to the Play Store. Tone: friendly and personal, not corporate."],
    "", [],
    "Deliverable:",
    ["Email copy in plain text, ready to send from any email client.",
     "Optional: HTML email template."]
)

add_task(doc, "E3", "Press Kit",
    ["Produce a one-page press kit PDF that Peter can send to bloggers, local press, or anyone "
     "interested in covering the app."],
    "Must include:",
    ["App name and logo/icon",
     "Short description (what it does, who it's for)",
     "2â€"3 key screenshots",
     "Developer contact details",
     "Play Store link"],
    "Deliverable:",
    ["restiview_press_kit.pdf â€" single page, suitable for both print and screen."]
)

# â"€â"€ Section 3: Out of Scope â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "3. Out of Scope")
items_oos = [
    "App development or code changes (developer responsibility)",
    "Play Store or Firebase account management",
    "Social media account creation or ongoing management",
    "Paid advertising or marketing campaigns",
    "Apple App Store assets (to be addressed in a future phase)",
]
for item in items_oos:
    bullet(doc, item)

# â"€â"€ Section 4: Deliverables Summary â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "4. Deliverables Summary")
add_table(doc,
    ["Ref", "Deliverable", "Format"],
    [
        ["A1", "Design audit document",           "PDF / annotated screenshots"],
        ["A2", "Colour & typography specification","PDF / Figma"],
        ["A3", "Design sign-off",                 "Written confirmation"],
        ["A4", "Component style specification",   "PDF / annotated screenshots"],
        ["B1", "Bug & feedback log",              "Spreadsheet / document"],
        ["B2", "Updated bug log",                 "Spreadsheet / document"],
        ["C1", "App icon files",                  "PNG + source"],
        ["C2", "Feature graphic",                 "PNG + source"],
        ["C3", "Short description options",       "Plain text"],
        ["C4", "Full Play Store description",     "Plain text"],
        ["D1", "Website audit",                   "Written list"],
        ["D2", "Updated live website",            "Published on restiview.com"],
        ["E1", "Social media graphics",           "PNG + source"],
        ["E2", "Launch email",                    "Text / HTML"],
        ["E3", "Press kit",                       "PDF"],
    ]
)

# â"€â"€ Section 5: Estimated Effort â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "5. Estimated Effort")
add_table(doc,
    ["Area", "Estimated Hours"],
    [
        ["A â€" App design review & improvements", "18h"],
        ["B â€" App testing",                      "7h"],
        ["C â€" Play Store assets",                "8h"],
        ["D â€" Website update",                   "6h"],
        ["E â€" Launch materials",                 "7h"],
        ["Total",                                "46h"],
    ]
)

# â"€â"€ Section 6: Priority Order â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "6. Suggested Priority Order")
priorities = [
    "C1, C2 â€" App icon and feature graphic  (needed for Play Store listing immediately)",
    "C3, C4 â€" Store descriptions  (needed for Play Store listing)",
    "A1, A2 â€" Design audit and proposal  (feeds into A3 and A4)",
    "B1      â€" App testing  (can run concurrently with design work)",
    "D2      â€" Website update  (once app is live)",
    "E1, E2, E3 â€" Launch materials  (for launch day)",
]
for i, p in enumerate(priorities, 1):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(f"{i}.  ")
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GREEN
    run2 = para.add_run(p)
    run2.font.size = Pt(10.5)

# â"€â"€ Section 7: Further Tasks â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "7. Further Tasks to Consider")
body(doc,
    "The following items are not in scope for this engagement but may be worth addressing "
    "before or after launch:")
further = [
    "Onboarding screens â€" A 2â€"3 slide welcome/tutorial shown on first launch explaining what "
    "the app does, helping new users get started without confusion.",
    "App preview video â€" A 15â€"30 second screen recording with captions showing the app in use; "
    "significantly boosts Play Store conversion rates.",
    "Apple App Store assets â€" Equivalent icons, screenshots (different sizes), and descriptions "
    "for the iOS release (future phase).",
    "Support / FAQ page â€" Both app stores ask for a support URL; a simple FAQ page on the "
    "website would satisfy this and reduce support queries.",
    "Social media presence â€" A dedicated Facebook or Instagram page for RestiView to direct "
    "users to for updates and support.",
    "App Store Optimisation (ASO) â€" Ongoing keyword research and description tuning to improve "
    "Play Store search ranking after launch.",
]
for item in further:
    bullet(doc, item)

# â"€â"€ Section 8: Compensation â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
h1(doc, "8. Compensation")
body(doc, "To be agreed separately between Peter Kern and Tara Kern.")

# â"€â"€ Footer â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Document prepared May 2026  Â·  RestiView  Â·  restiview.com")
run.font.size = Pt(9)
run.font.color.rgb = MID_GREY

# â"€â"€ Save â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
doc.save(OUTPUT_PATH)
print(f"Saved: {os.path.abspath(OUTPUT_PATH)}")

