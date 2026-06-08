"""Generate a publishable RestiView user guide as a Word document.

Run from the project root with:
    .venv\\Scripts\\python tool\\generate_user_guide.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "RestiView_User_Guide.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_text(doc, text):
    for block in text.split("\n\n"):
        doc.add_paragraph(block)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("RestiView User Guide")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(32, 74, 52)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "A practical guide for creating, managing, and sharing restaurant reviews"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11.5)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run(
        "Website: www.restiview.com\nCurrent app guide edition: June 2026"
    ).font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading("1. Introduction", level=1)
    add_text(
        doc,
        "RestiView is a personal restaurant review app. It helps you build your own library of memorable dining experiences, search that library later, and share selected reviews with people you trust. Unlike public review platforms, RestiView is designed around your own records and your own network of friends and family.",
    )
    add_bullets(
        doc,
        [
            "Create and save restaurant reviews for places you have visited.",
            "Search and filter your review history by location, cuisine, rating, and occasion.",
            "Add friends and exchange selected reviews privately.",
            "Request reviews from friends when planning a trip or meal.",
            "Customise cuisines, occasions, and countries to match your own style.",
        ],
    )

    doc.add_heading("2. Before You Start", level=1)
    add_bullets(
        doc,
        [
            "You need an internet connection to register, sign in, sync data, and exchange reviews.",
            "Location permission is optional, but it improves restaurant auto-fill when adding a review.",
            "Photo permission is optional, but it is required if you want to attach photos to review detail cards.",
            "Friends features are optional and can be enabled or disabled in Settings.",
        ],
    )

    doc.add_heading("3. Getting Started", level=1)
    doc.add_heading("Create an account", level=2)
    add_numbered(
        doc,
        [
            "Open RestiView and tap REGISTER.",
            "Enter your email address, password, display name, and home country.",
            "Review and accept the Terms and Conditions.",
            "Complete registration to create your account.",
        ],
    )
    add_text(
        doc,
        "Your home country is used as a default when you create reviews and helps the app apply sensible defaults later.",
    )

    doc.add_heading("Sign in", level=2)
    add_numbered(
        doc,
        [
            "Tap SIGN IN.",
            "Enter the email address and password you registered with.",
            "Optionally enable Stay Signed In if you want RestiView to remember your session on that device.",
        ],
    )

    doc.add_heading("Reset a forgotten password", level=2)
    add_numbered(
        doc,
        [
            "On the sign-in screen, tap Forgot Password?.",
            "Enter your account email address.",
            "Use the reset link sent to your email inbox.",
        ],
    )

    doc.add_heading("4. Home Screen Overview", level=1)
    add_text(
        doc,
        "After signing in, RestiView opens to the main dashboard. The exact layout may vary slightly by device, but the key actions remain the same.",
    )
    add_bullets(
        doc,
        [
            "ADD REVIEW: start a new restaurant review.",
            "VIEW REVIEWS: open your saved review library.",
            "REQUESTED REVIEWS: view reviews delivered to you by friends after a review request.",
            "FRIENDS: manage friend relationships and review requests.",
            "SETTINGS: change defaults, permissions, and account options.",
            "HELP: open the in-app guide and support links.",
            "SIGN OUT: end the current session on the device.",
        ],
    )
    add_text(
        doc,
        "Notification badges may appear on the Friends area when there are incoming friend requests or review requests waiting for attention.",
    )

    doc.add_heading("5. Adding a Review", level=1)
    add_text(
        doc,
        "To add a review, tap ADD REVIEW on the home screen. The review wizard is divided into four stages.",
    )

    doc.add_heading("Step 1: General Info", level=2)
    add_bullets(
        doc,
        [
            "Restaurant name",
            "City",
            "Country",
            "Cuisine",
            "Occasion",
            "Visit date",
            "Number of diners",
            "Cost",
        ],
    )
    add_text(
        doc,
        "If Location Services are enabled, RestiView can search nearby restaurants and auto-fill restaurant details for you. If location is unavailable or permission is denied, you can always enter the details manually.",
    )

    doc.add_heading("Step 2: Comments and Details", level=2)
    add_bullets(
        doc,
        [
            "Add your free-text comments about the visit.",
            "Record optional meal details such as starters, mains, desserts, wine, cocktails, or drinks.",
            "Attach photos to detail cards if photo access is enabled in Settings.",
        ],
    )

    doc.add_heading("Step 3: Ratings", level=2)
    add_bullets(
        doc,
        [
            "Food",
            "Service",
            "Ambiance",
            "Drinks",
            "Value For Money",
            "Michelin stars, where relevant",
        ],
    )
    add_text(
        doc,
        "Ratings are entered on a 0 to 5 star basis. RestiView uses these values later for filtering and sorting.",
    )

    doc.add_heading("Step 4: Good For", level=2)
    add_text(
        doc,
        "Tag the restaurant with one or more occasion-style categories such as Date Night, Business Lunch, Brunch, or Large Groups. These tags make it easier to search your reviews later.",
    )

    doc.add_heading("Preview and save", level=2)
    add_numbered(
        doc,
        [
            "Tap PREVIEW to review the full entry before saving.",
            "Check the restaurant details, comments, ratings, and tags.",
            "Tap SAVE to store the review permanently.",
        ],
    )
    add_text(
        doc,
        "If you leave a review part way through, RestiView may offer to resume a locally saved draft the next time you open the app.",
    )

    doc.add_heading("6. Viewing, Searching, and Editing Reviews", level=1)
    add_text(doc, "Tap VIEW REVIEWS to open your personal review library.")
    add_bullets(
        doc,
        [
            "Sort by Date, Rating, Name, City, or Cuisine.",
            "Filter by Country, City, or Cuisine.",
            "Apply a minimum star rating threshold.",
            "Use Good For tags to narrow results to specific occasions.",
        ],
    )
    add_text(
        doc,
        "Tap any review card to open the full preview. From there you can edit the review, update the content, or delete it if you no longer want it in your library.",
    )

    doc.add_heading("7. Friends", level=1)
    add_text(
        doc,
        "The Friends feature lets you connect privately with other RestiView users so you can exchange reviews directly.",
    )
    doc.add_heading("Send a friend request", level=2)
    add_numbered(
        doc,
        [
            "Tap FRIENDS on the home screen.",
            "Tap +FRIEND.",
            "Enter the other user's email address.",
            "Submit the request.",
        ],
    )
    add_text(
        doc,
        "The other user must have Allow Friends enabled in their Settings before you can send them a request.",
    )
    doc.add_heading("Respond to an incoming friend request", level=2)
    add_numbered(
        doc,
        [
            "Look for the notification marker on the FRIENDS area.",
            "Open FRIENDS and select the incoming request row.",
            "Choose ACCEPT or DECLINE.",
        ],
    )

    doc.add_heading("8. Review Requests", level=1)
    add_text(
        doc,
        "Review Requests are useful when you trust another user's restaurant knowledge in a city or country and want a curated copy of their reviews.",
    )
    doc.add_heading("Send a review request", level=2)
    add_numbered(
        doc,
        [
            "Open FRIENDS and select one of your accepted friends.",
            "Review the countries and cities where that friend has reviews.",
            "Tick the locations you want and tap REQUEST.",
            "Optionally include a short request comment.",
        ],
    )
    doc.add_heading("Provide reviews when you receive a request", level=2)
    add_numbered(
        doc,
        [
            "Open FRIENDS when the request notification appears.",
            "Open the request details to see who asked, what they asked for, and how many of your reviews match.",
            "Tap REVIEW to inspect the matching reviews one by one.",
            "Choose whether to include or exclude individual reviews.",
            "Tap ACCEPT to send the approved set, or DECLINE to refuse the request.",
        ],
    )
    add_text(
        doc,
        "Delivered reviews appear under REQUESTED REVIEWS on the requester's home screen. These are separate from your own authored reviews and help you build a trusted private recommendation list from friends.",
    )

    doc.add_heading("9. Settings", level=1)
    add_text(
        doc,
        "Settings control both your personal defaults and several optional permissions-based features.",
    )
    settings_table = doc.add_table(rows=1, cols=2)
    settings_table.style = "Table Grid"
    hdr_cells = settings_table.rows[0].cells
    set_cell_text(hdr_cells[0], "Setting", bold=True)
    set_cell_text(hdr_cells[1], "What it does", bold=True)
    shade_cell(hdr_cells[0], "D9EAD3")
    shade_cell(hdr_cells[1], "D9EAD3")
    settings_rows = [
        ("Sort By", "Sets the default sort order for your review list."),
        ("Home Country", "Defines the default country used when creating reviews."),
        ("Allow Location Services", "Enables restaurant auto-fill and nearby place detection while adding reviews."),
        ("Search Radius", "Controls how far the nearby restaurant search looks, typically between 10 and 200 metres."),
        ("Allow Photos", "Enables photo attachments on review detail cards."),
        ("Allow Friends", "Controls whether other users can send you friend requests or review requests."),
        ("Allow Auto Capture", "Reserved for future functionality. It is currently inactive."),
        ("Custom Values", "Lets you add or maintain your own cuisines, occasions, and countries for dropdown use."),
        ("Reset Settings", "Restores settings to their default values."),
        ("Save Changes", "Writes any settings updates to your account. Changes are not saved until you tap this."),
        ("Delete Account", "Permanently deletes your account and associated data. This cannot be undone."),
    ]
    for left, right in settings_rows:
        row_cells = settings_table.add_row().cells
        set_cell_text(row_cells[0], left, bold=True)
        set_cell_text(row_cells[1], right)

    doc.add_heading("10. Custom Values", level=1)
    add_text(
        doc,
        "If the built-in dropdown choices are not enough, you can create your own cuisine types, occasion tags, and countries. These custom values then appear in relevant dropdowns throughout the app. This is useful if you travel widely or use your own personal naming conventions.",
    )

    doc.add_heading("11. Permissions and Privacy", level=1)
    add_bullets(
        doc,
        [
            "Location is used to help find nearby restaurants when creating a review.",
            "Photos are used only when you choose to attach images to review content.",
            "Friends and review exchanges are private app-to-app features and are not public social posts.",
            "You remain in control of whether to accept friend requests, respond to review requests, and include or exclude individual reviews before sharing.",
        ],
    )

    doc.add_heading("12. Troubleshooting", level=1)
    troubleshooting = [
        (
            "I cannot find a nearby restaurant automatically.",
            "Check that Allow Location Services is enabled, your phone location is turned on, and your Search Radius is sensible. If auto-fill still does not find the venue, enter the details manually.",
        ),
        (
            "I cannot add photos.",
            "Make sure Allow Photos is enabled in Settings and that the device has granted photo access to RestiView.",
        ),
        (
            "I cannot send a friend request.",
            "Verify the email address is correct and that the other user has Allow Friends enabled.",
        ),
        (
            "I requested reviews but nothing arrived yet.",
            "The other user must open the request, review the matching items, and accept the delivery before the reviews appear under Requested Reviews.",
        ),
        (
            "I forgot my password.",
            "Use Forgot Password? from the sign-in screen to send a reset link to your registered email address.",
        ),
    ]
    for question, answer in troubleshooting:
        doc.add_heading(question, level=3)
        doc.add_paragraph(answer)

    doc.add_heading("13. Practical Tips", level=1)
    add_bullets(
        doc,
        [
            "Add Good For tags consistently so your searches stay useful over time.",
            "Use Custom Values sparingly and consistently to avoid duplicate categories with slightly different names.",
            "Review your preview screen before saving so ratings, cost, and tags are complete.",
            "Only share reviews you are comfortable sending to the recipient. Use the include/exclude controls carefully.",
            "Keep Home Country accurate so the app can warn you when your current location differs from your usual default.",
        ],
    )

    doc.add_heading("14. Support", level=1)
    add_text(
        doc,
        "For more information, updates, and website publication of this guide, visit www.restiview.com. In-app Help also provides a quick version of the same workflows for day-to-day use.",
    )

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("RestiView User Guide")
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUTPUT)
    print(f"Created: {OUTPUT}")