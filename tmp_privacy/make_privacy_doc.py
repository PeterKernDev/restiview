from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import docx.oxml as oxml

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Styles helper ───────────────────────────────────────────────────────────
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Calibri'

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

# ── Document content ─────────────────────────────────────────────────────────

add_title('RestiView Privacy Policy')
add_subtitle('Last updated: 10 April 2026')
doc.add_paragraph()

add_body(
    'This privacy policy applies between you, the User of this App and Website, and RestiView, '
    'the owner and provider of this App and Website. RestiView takes the privacy of your information '
    'very seriously. This policy outlines how we collect, use, and protect your data in accordance '
    'with applicable data protection laws, including the UK GDPR.'
)

# 1
add_heading('1. Definitions')
add_bullet(' All personal information submitted by you or collected automatically during your use of the App or Website.', 'Data:')
add_bullet(' Includes the UK GDPR and any applicable national legislation.', 'Data Protection Laws:')
add_bullet(' RestiView, based in Newmarket, Suffolk, UK.', 'RestiView, we, or us:')
add_bullet(' Any individual accessing the App or Website who is not employed by RestiView or providing services to it.', 'User or you:')
add_bullet(' The RestiView mobile application and www.restiview.com, including subdomains.', 'App and Website:')

# 2
add_heading('2. Scope')
add_body('This policy applies only to data collected by RestiView through the App and Website. It does not apply to third-party services or external links.')

# 3
add_heading('3. Data We Collect')
add_body('We may collect the following types of personal data:')
add_bullet('Name')
add_bullet('Email address and contact details')
add_bullet('Location data (if permission is granted)')
add_bullet('Images captured within the app — stored locally on your device and not transmitted to RestiView servers')
add_bullet('Device information and usage statistics')
add_bullet('Authentication and account data via Firebase')
add_bullet('Display name and email address, which may be visible to other users you connect with via the Friends feature')

# 4
add_heading('4. How We Collect Data')
add_bullet(' When you register, contact us, or upload content', 'Provided by you:')
add_bullet(' Through app usage, analytics, and device interactions', 'Automatically:')
add_bullet(' Firebase Authentication, Firebase Realtime Database, Google Play Services, and Google Places API (used to find nearby restaurants based on your location)', 'Third-party services:')

# 5
add_heading('5. Use of Data')
add_body('We use your data to:')
add_bullet('Provide and improve app functionality')
add_bullet('Authenticate users and manage accounts')
add_bullet('Store and retrieve user-generated content')
add_bullet('Suggest nearby restaurants based on your location')
add_bullet('Enable the Friends feature, allowing users to share restaurant reviews with each other')
add_bullet('Respond to support requests')
add_bullet('Analyse usage patterns to improve performance')
add_bullet('Comply with legal obligations')

# 6
add_heading('6. Legal Basis for Processing')
add_body('We process your data under the following legal bases:')
add_bullet(' For optional features like location and media access', 'Consent:')
add_bullet(' To deliver services you\'ve requested', 'Contract:')
add_bullet(' For analytics, security, and service improvement', 'Legitimate interests:')
add_bullet(' Where required by law', 'Legal compliance:')

# 7
add_heading('7. Data Security')
add_body('We implement technical and organisational measures to protect your data:')
add_bullet('Secure servers and encrypted connections')
add_bullet('Firebase security rules and access controls')
add_bullet('Password-protected user accounts')
add_bullet('Monitoring for unauthorised access or breaches')
add_body(
    'Your data may be stored and processed outside the UK or EEA, including in the United States, '
    'via Firebase (Google LLC). Google LLC participates in applicable international data transfer '
    'frameworks. By using the app you consent to this transfer.'
)
add_body('If you suspect misuse or a breach, contact us immediately at RestiView@gmail.com.')

# 8
add_heading('8. Data Retention')
add_body(
    'We retain your data only as long as necessary for the purposes outlined in this policy or as '
    'required by law. Backup data may persist for legal or regulatory reasons. You may request '
    'deletion of your data at any time (see Your Rights below).'
)

# 9
add_heading('9. Your Rights')
add_body('You have the right to:')
add_bullet('Access, correct, or delete your data')
add_bullet('Restrict or object to processing')
add_bullet('Request data portability')
add_bullet('Withdraw consent at any time')
add_bullet('Delete your account directly within the app via Settings → Delete Account')
add_body(
    'To exercise these rights, email RestiView@gmail.com. If unsatisfied, you may contact the '
    'UK Information Commissioner\'s Office (ICO) at https://ico.org.uk.'
)

# 10
add_heading('10. Children\'s Privacy')
add_body(
    'This app is not directed at children under the age of 13. We do not knowingly collect personal '
    'data from children. If you believe a child has provided us with personal data, please contact '
    'us so we can delete it.'
)

# 11
add_heading('11. Third-Party Links')
add_body(
    'Our App and Website may contain links to third-party services. We are not responsible for their '
    'privacy practices. Please review their policies before use.'
)

# 12
add_heading('12. Business Transfers')
add_body(
    'If RestiView is sold or merged, your data may be transferred to the new owner under the terms '
    'of this policy.'
)

# 13
add_heading('13. Changes to This Policy')
add_body(
    'We may update this policy from time to time. Changes will be posted on our Website and reflected '
    'in the App. Continued use indicates acceptance of the revised terms.'
)

# 14
add_heading('14. Contact Us')
add_body('For questions, concerns, or data requests, contact:')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
p.add_run('RestiView\n').font.name = 'Calibri'
p.add_run('Newmarket, Suffolk, UK\n').font.name = 'Calibri'
p.add_run('RestiView@gmail.com').font.name = 'Calibri'
for run in p.runs:
    run.font.size = Pt(11)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'c:\dev\RestiView2\restiview\tmp_privacy\RestiView_Privacy_Policy.docx'
doc.save(out)
print(f'Saved: {out}')
